from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
document.sections[0].page_width = Inches(10)
document.sections[0].left_margin=Inches(1/3)
document.sections[0].right_margin=Inches(1/3)
table = document.add_table(rows=1, cols=3)
row = table.rows[0]
cell1 = row.cells[0]
image_path = "C:\\Users\\b8708\\OneDrive\\Рабочий стол\\head.png"
cell1_paragraph = cell1.paragraphs[0]
run = cell1_paragraph.add_run()
run.add_picture(image_path, width=Inches(2.0))
cell2 = row.cells[2]
text_content = "УТВЕРЖДАЮ\nДекан Школы/Центра\nФИO\n_________________\n«___» ______202____"
cell2.text = text_content


table = document.add_table(rows=4, cols=3)
table.style = 'Table Grid'
row = table.rows[0]
row.cells[0].text = "КОД И НАЗВАНИЕ ДИСЦИПЛИНЫ:\nFA1307 Research Methods"

row.cells[1].text = "КРЕДИТЫ ECTS И ЧАСЫ:\n3 ECTS\nВсего часов: 90\nАудиторные часы: 30 часов\nСамостоятельная работа\n(СРОП, СРО): 60 часов"

cell13 = row.cells[2]
text_content = "ПРЕРЕКВИЗИТЫ:\nНет"
cell13.text = text_content

row = table.rows[1]
cell11 = row.cells[0]
text_content = "УРОВЕНЬ ОБУЧЕНИЯ:\nБакалавриат"
cell11.text = text_content

cell12 = row.cells[1]
text_content = "СЕМЕСТР:\n7"
cell12.text = text_content

cell13 = row.cells[2]
text_content = "ОБРАЗОВАТЕЛЬНАЯ ПРОГРАММА: \nФинансы"
cell13.text = text_content

row = table.rows[2]
cell11 = row.cells[0]
text_content = "ЯЗЫК ОБУЧЕНИЯ:\nРусский"
cell11.text = text_content

cell12 = row.cells[1]
text_content = "УРОВЕНЬ ВЛАДЕНИЯ ЯЗЫКОМ ОБУЧЕНИЯ:\nA1"
cell12.text = text_content

cell13 = row.cells[2]
text_content = "ФОРМАТ ОБУЧЕНИЯ:\nOffline"
cell13.text = text_content

row = table.rows[3]
cell11 = row.cells[0]
text_content = "ПРЕПОДАВАТЕЛЬ:\nФИО, должность, ученая степень"
cell11.text = text_content

cell12 = row.cells[1]
text_content = "КОНТАКТЫ ПРЕПОДАВАТЕЛЯ:\nЭл.почта/ телефон: x.xxxxx@almau.edu.kz"
cell12.text = text_content

cell13 = row.cells[2]
text_content = "ВРЕМЯ И МЕСТО ПРОВЕДЕНИЯ ЗАНЯТИЙ:\nПо утвержденному расписанию"
cell13.text = text_content

text_to_write = "Цель курса\nДанная дисциплина нацелена на поддержание и развитие аналитического, критического мышления и творческих навыков, а также написания и презентации исследования."
paragraph = document.add_paragraph(text_to_write)

text_to_write = "График занятий и задания"
paragraph = document.add_paragraph(text_to_write)

table1 = document.add_table(rows=9, cols=4)
table1.style = 'Table Grid'
column = table1.columns[0]
column.width = Inches(0.4)
table1.rows[0].cells[0].text='Недели'
table1.rows[0].cells[1].text='Тема / модуль'
table1.rows[0].cells[2].text='Формат проведения занятий'
table1.rows[0].cells[3].text='Задания'

for i in range(1, 9):
    table1.rows[i].cells[0].text = str(i+1) + "-" + str(i+2)
    table1.rows[i].cells[0].width = Inches(0.4)


paragraph = document.add_paragraph('\n\n\n')
new_table1 = document.add_table(rows=1, cols=2)

row = new_table1.rows[0]
table1 = row.cells[0].add_table(rows=4, cols=1)
table1.rows[0].cells[0].text='Academic Handbook'
run = table1.rows[1].cells[0].paragraphs[0].add_run()
run.add_picture("C:\\Users\\b8708\\OneDrive\\Изображения\\1.gif", width=Inches(1.0))
table1.rows[2].cells[0].text='https://almauedu-my.sharepoint.com/:f:/g/personal/f_abdoldina_almau_edu_kz/EnVy7hCS47hMoVtpgjfq3-YBY2biThYahFoceoI9xY1n3A?e=wASl1u'
table1.rows[3].cells[0].text='Составлено:\nк.э.н., lecturer	___________	ФИО ППС'
table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


table1 = row.cells[1].add_table(rows=4, cols=1)
table1.rows[0].cells[0].text='Результаты обучения курса. Задания и политика курса:'
run = table1.rows[1].cells[0].paragraphs[0].add_run()
run.add_picture("C:\\Users\\b8708\\OneDrive\\Изображения\\1.gif", width=Inches(1.0))
table1.rows[2].cells[0].text='https://docs.google.com/document/u/0/d/19QyuM6a1uyAXd49Rb9cpOW43lyYe5Cc0/mobilebasic'
table1.rows[3].cells[0].text='Согласовано:\nДиректор УМ	___________	Абдолдина Ф.Н.'
table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('\n\n\n')
document.add_paragraph('1.	Описание курса')
document.add_paragraph('	Дисциплина «Research Methods» предлагает студентам всестороннее представление о том, как проводить исследовательскую работу/проект и профессионально презентовать достигнутые результаты, учитывая специфику предстоящих научных исследований по специальностям, и формирует компетенции в области научных исследований. Переходя от первых шагов исследовательской работы/проекта (определение проблемных вопросов исследования) к последнему результату (разработка рекомендаций), в конце курса студенты смогут провести углубленное исследование в своей области, и смогут более уверенно представлять и защищать свои идеи перед критической аудиторией.'+
'\n	Основная цель курса - формирование креативного исследовательского мышления и способностей решать разнообразные хозяйственные, социальные, психологические задачи путем использования современных методов, приемов и средств научного исследования.'+
'\n	Теоретический подход сочетается с практическими заданиями по проводимым исследованиям и выбранной темы дипломного проекта.')
document.add_paragraph('\n\n')
document.add_paragraph('2.	Таблица соответствия Результатов обучения курса Результатам обучения образовательной программы')

table = document.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Результаты обучения курса (РО курса)'
table.rows[0].cells[1].text = 'Результаты обучения образовательной программы (РО ОП)'
table.rows[1].cells[0].text = 'Теоретические и практические знания\n'+'РО 1. Определять основные понятия и отрывки для планирования, разработки и проведения исследований\n'+'РО 2. Описывать фундаментальные этические стандарты научных исследований\n''РО 3. Определять различные методологии\n'+'РО 4. Объяснять и применять различные методы исследования (качественные и количественные)\n'+'РО 5. Планировать и структурировать исследования в своей области исследований\n'
table.rows[1].cells[1].text='ON5 описывать и применять основные методы и инструменты научного исследования, владение математическими и экономико-статистическими и финансовыми методами и инструментами для подготовки и проведения финансового анализа и оценки эффективности операционной, финансовой и инвестиционной деятельности компании;'
table.rows[2].cells[0].text='Когнитивные и практические навыки и компетенции\n'+'РО 6. Разработка планов исследований, работающих индивидуально и / или в группе\n'+'РО 7. Применение различных методов сбора данных (библиографические исследования, исследования в режиме онлайн, интервью, опросы)\n'+'РО 8. Применение различных качественных и количественных подходов к анализу данных;\n'+'РО 9. Организация в эффективном представлении информации (защита планов исследования)\n'
table.rows[2].cells[1].text='ON10 применять базовые исследовательские навыки, информационные и финансовые технологии, навыки критического мышления, коммуникационные навыки для выбора подходящих теорий и методологий, получения актуальной и точной информации, анализа данных и разработки выводов в теоретических исследованиях или прикладных проектах.'

document.add_paragraph('\n\n3.	Тематический план')
table = document.add_table(rows=1,cols=7)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '1-2'
table.rows[0].cells[1].text = 'Тема / модуль'
table.rows[0].cells[2].text = 'РО курса, РО ОП'
table.rows[0].cells[3].text = 'Вопросы по теме / модулю'
table.rows[0].cells[4].text = 'Задания'
table.rows[0].cells[5].text = 'Литература'
table.rows[0].cells[6].text = 'Структура оценок'


document.add_paragraph('\n\n4.	Система оценивания курса')
table = document.add_table(rows=1,cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Тема / модуль'
table.rows[0].cells[1].text = 'Максимальный процент(%)'
table.rows[0].cells[2].text = 'Максимальный вес(%)'
table.rows[0].cells[3].text = 'Итого в баллах'


document.add_paragraph('\n\n5.	Список литературы')

document.add_paragraph("""Обязательная литература
    1.	Новиков, А.М. Методология научного исследования [Текст]: учебно-методическое пособие/ А.М. Новиков, Д.А. Новиков. – Изд. 2-е, - Москва: Кн.дом “ЛИБРОКОМ”:URSS, 2013. -270 с. 
    2.	An introduction to Business research methods Dr. Sue Greener, Dr. Joe Martell.- 2nd. ed.- Bookboon. com., 2015.- 137 p. 
    3.	Герасимов Б.И. Основы научных исследований [Электронный ресурс] / Б.И. Герасимов, В.В. Дробышева, Н.В. Злобина и др. - М.: Форум: НИЦ Инфра-М, 2013. Режим доступа: http://znanium.com/bookread.php?book=390595 (дата обращения 02.09.2016)

Дополнительная литература 
4.	Robert, K.Yin.Case Study Research [Текст]: Design and Methods / K.Yin Robert.- USA: Sage, 2014.- 282 с.
5.	Орехов А.М. Методы экономических исследований [Электронный ресурс]: Учебное пособие / А.М. Орехов. - 2-e изд. - М.: НИЦ Инфра-М, 2013. - 344 с. Режим доступа: http://znanium.com/bookread.php?book=362627 (дата обращения 15.03.2015) 
6.	Мильчакова, Н.Н., Яркова, Е.Н. Методы социально-экономических исследований: учебное пособие/ Н.Н. Мильчакова, Е. Н. Яркова; Тюм. гос. ун-т. - Тюмень: Изд-во ТюмГУ, 2014. - 379 с.

Интернет ресурсы
    1.	Библиотека AlmaU  http://lib.almau.edu.kz/
    2.	Научная электронная библиотека http://elibrary.ru/ 
    3.	Научно-образовательный портал: http://www.med-edu.ru/
    4.	Международные организации Организация Объединенных Наций(ООН)- United Nations(UN)- http://www.un.org/ 
    5.	Международный валютный фонд(МВФ)- International Monetary Fund – IMFhttp://www.imf.org Всемирный Банк (World Bank)- http://www.worldbank.org 
    6.	Всемирная организация интеллектуальной собственности (ВОИС)- World Intellectual Property Organization (WIPO) - http://www.wipo.org 
    7.	Всемирный экономический форум - World Economic Forum- http://www.weforum.org БРИКС http://infobrics.org ШОС http://infoshos.ru 
    8.	Национальный банк РК. Официальный интернет ресурс http://www.nationalbank.kz/?switch=russian
    9.	Казахстанская фондовая биржа (KASE) Официальный интернет ресурс kase.kz
    10.	Кафедра экономической методологии и истории Высшей школы экономики: курсы, публикации http://www.hse.ru/kafedry/economy/ec_methodology_history/default.htm 
    11.	Портал по социологии, экономике и менеджменту www.ecsocman.edu.ru 
    12.	Портал по общественным наукам www.socionet.ru 

""")

document.add_paragraph("""6.	Философия преподавания и обучения
    Процесс обучения основывается на освоении теоретического материала на лекциях, на самостоятельном изучении материалов, практического применения знаний и обсуждениях в аудитории. Студенты, обучаясь в условиях использования активных форм, работая в группах, решая конкретные ситуационные задачи, приобретут способность при¬нимать решения в нестандартных ситуациях, умение работать в команде, самостоятельно добывать, анализировать и эффективно использовать информацию, рационально работать.
    Задача преподавателя будет заключаться в том, чтобы обеспечить учебным материалом, рекомендуемой литературой, донести сложные аспекты в доступной форме. Преподаватель несет ответственность за успешное освоение знаний и навыков в течение контактных часов и в процессе руководства самостоятельной работой студентов
    Подведение итогов преподавателем в конце недели позволяет студентам видеть свои еженедельные результаты, образующие средневзвешенные оценки уровня достижений (GPI).
""")

document.add_paragraph("""7.	Политика курса
Этика занятий
Освоение дисциплины «Research Methods» предусматривает 
-	обязательное посещение занятий;
-	активность во время занятий;
-	подготовка к занятиям, выполнение домашнего задания;
-	сдача заданий в установленные сроки;
-	быть терпимым, открытым и доброжелательным; 
-	конструктивно поддерживать обратную связь на всех занятиях; 
-	быть пунктуальным и обязательным.
Недопустимо:
-	пропуски по неуважительным причинам;
-	опоздание и уход с занятий (В случае опоздания студент не допускается на занятие, т.к. он нарушает ход учебного занятия);
-	несвоевременная сдача заданий и др. 
При пропусках занятий по уважительной причине допускается отработка пройденного материала.

Этика экзамена
Недопустимо:
-	опоздание; 
-	пользование мобильными телефонами во время экзамена; 
-	списывание при сдаче экзамена. За списывание на контрольном мероприятии студент удаляется из аудитории и ему выставляется 0 баллов.
Если в силу каких-либо уважительных причин вы отсутствовали во время проведения контрольного мероприятия, вам предоставляется возможность пройти его в дополнительно назначенное преподавателем время (РК и ИК сдаются с разрешения декана), в противном случае вы получаете «0» баллов.
Политика академического поведения и этики основана на Кодексе корпоративной культуры, Этическом кодексе студента, Правилах внутреннего распорядка AlmaU.

Информация и связь
Вы должны регулярно (ежедневно) проверять Личную страницу в автоматизированной информационной системе, LMS и электронную почту, чтобы получать дополнительную информацию, задания или знать изменения в расписании. 

""")


document.save('demo.docx')



