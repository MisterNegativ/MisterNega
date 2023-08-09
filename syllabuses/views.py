from django.shortcuts import redirect, render
from django.shortcuts import render
from django.urls import reverse
from syllabuses.models import *
from .forms import *
from django.views import View
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect, JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.http import FileResponse
from docx import Document
from reportlab.pdfgen import canvas
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import win32com.client as client
import pythoncom
from django.http import HttpResponse




def download_syllabus_as_pdf_rus(request, syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    document = Document()
    document.sections[0].page_width = Inches(10)
    document.sections[0].left_margin = Inches(2)
    document.sections[0].right_margin = Inches(2)
    table = document.add_table(rows=1, cols=3)
    row = table.rows[0]
    cell1 = row.cells[0]
    image_path = "C:\\Users\\User\\Desktop\\gorizontal'niy rgb.png"
    cell1_paragraph = cell1.paragraphs[0]
    run = cell1_paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.0))
    cell2 = row.cells[2]
    text_content = "УТВЕРЖДАЮ\nДекан Школы/Центра\nФИO\n_________________\n«___» ______202____"
    cell2.text = text_content

    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    row = table.rows[0]
    row.cells[0].text = "КОД И НАЗВАНИЕ ДИСЦИПЛИНЫ:\n" + syllabus.syllabus_name + " " + syllabus.course.__str__()
    hours = syllabus.total_hours-syllabus.classroom_hours
    row.cells[1].text = "КРЕДИТЫ ECTS И ЧАСЫ:\n3 ECTS\nВсего часов:" + str(syllabus.total_hours) + "\nАудиторные часы:" + str(syllabus.classroom_hours) + "часов\nСамостоятельная работа\n(СРОП, СРО):" +  str(hours) + "часов"

    cell13 = row.cells[2]
    text_content = "ПРЕРЕКВИЗИТЫ:\n" + syllabus.prerequisites
    cell13.text = text_content

    row = table.rows[1]
    cell11 = row.cells[0]
    text_content = "УРОВЕНЬ ОБУЧЕНИЯ:\n" + str(syllabus.training_level)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "СЕМЕСТР:\n" + str(syllabus.semester)
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ОБРАЗОВАТЕЛЬНАЯ ПРОГРАММА: \n" + str(syllabus.edu_programms)
    cell13.text = text_content

    row = table.rows[2]
    cell11 = row.cells[0]
    text_content = "ЯЗЫК ОБУЧЕНИЯ:\n" + str(syllabus.language_of_education)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "УРОВЕНЬ ВЛАДЕНИЯ ЯЗЫКОМ ОБУЧЕНИЯ:\n" + str(syllabus.proficiency_level)
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ФОРМАТ ОБУЧЕНИЯ:\n" + str(syllabus.format_of_training)
    cell13.text = text_content

    row = table.rows[3]
    cell11 = row.cells[0]
    text_content = "ПРЕПОДАВАТЕЛЬ:\n" + str(syllabus.instructor)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "КОНТАКТЫ ПРЕПОДАВАТЕЛЯ:\nЭл.почта/ телефон:" + syllabus.instructor.email
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ВРЕМЯ И МЕСТО ПРОВЕДЕНИЯ ЗАНЯТИЙ:\n" + syllabus.time_place
    cell13.text = text_content

    text_to_write = "Цель курса\nДанная дисциплина нацелена на поддержание и развитие аналитического, критического мышления и творческих навыков, а также написания и презентации исследования."
    paragraph = document.add_paragraph(text_to_write)

    text_to_write = "График занятий и задания"
    paragraph = document.add_paragraph(text_to_write)
    modules = Module.objects.filter(syllabus=syllabus)
    table1 = document.add_table(rows=len(modules)+1, cols=4)
    table1.style = 'Table Grid'
    column = table1.columns[0]
    column.width = Inches(0.4)
    table1.rows[0].cells[0].text = 'Недели'
    table1.rows[0].cells[1].text = 'Тема / модуль'
    table1.rows[0].cells[2].text = 'Формат проведения занятий'
    table1.rows[0].cells[3].text = 'Задания'
    j = 1
    for i in modules:
        table1.rows[j].cells[0].text = str(i.week)
        table1.rows[j].cells[0].width = Inches(0.4)
        table1.rows[j].cells[1].text = i.theme
        table1.rows[j].cells[2].text = str(i.format.__str__())
        table1.rows[j].cells[3].text = i.tasks
        j+=1
    document.add_paragraph('\n\n\n')
    new_table1 = document.add_table(rows=1, cols=2)

    row = new_table1.rows[0]
    table1 = row.cells[0].add_table(rows=4, cols=1)
    table1.rows[0].cells[0].text = 'Academic Handbook'
    run = table1.rows[1].cells[0].paragraphs[0].add_run()
    table1.rows[2].cells[
        0].text = 'https://almauedu-my.sharepoint.com/:f:/g/personal/f_abdoldina_almau_edu_kz/EnVy7hCS47hMoVtpgjfq3-YBY2biThYahFoceoI9xY1n3A?e=wASl1u'
    table1.rows[3].cells[0].text = 'Составлено:\nк.э.н., lecturer	___________	ФИО ППС'
    table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = row.cells[1].add_table(rows=4, cols=1)
    table1.rows[0].cells[0].text = 'Результаты обучения курса. Задания и политика курса:'
    run = table1.rows[1].cells[0].paragraphs[0].add_run()
    table1.rows[2].cells[
        0].text = 'https://docs.google.com/document/u/0/d/19QyuM6a1uyAXd49Rb9cpOW43lyYe5Cc0/mobilebasic'
    table1.rows[3].cells[0].text = 'Согласовано:\nДиректор УМ	___________	Абдолдина Ф.Н.'
    table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('\n\n\n')
    document.add_paragraph('1.	Описание курса')
    document.add_paragraph(
        '	Дисциплина «Research Methods» предлагает студентам всестороннее представление о том, как проводить исследовательскую работу/проект и профессионально презентовать достигнутые результаты, учитывая специфику предстоящих научных исследований по специальностям, и формирует компетенции в области научных исследований. Переходя от первых шагов исследовательской работы/проекта (определение проблемных вопросов исследования) к последнему результату (разработка рекомендаций), в конце курса студенты смогут провести углубленное исследование в своей области, и смогут более уверенно представлять и защищать свои идеи перед критической аудиторией.' +
        '\n	Основная цель курса - формирование креативного исследовательского мышления и способностей решать разнообразные хозяйственные, социальные, психологические задачи путем использования современных методов, приемов и средств научного исследования.' +
        '\n	Теоретический подход сочетается с практическими заданиями по проводимым исследованиям и выбранной темы дипломного проекта.')
    document.add_paragraph('\n\n')
    document.add_paragraph(
        '2.	Таблица соответствия Результатов обучения курса Результатам обучения образовательной программы')

    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Результаты обучения курса (РО курса)'
    table.rows[0].cells[1].text = 'Результаты обучения образовательной программы (РО ОП)'
    table.rows[1].cells[
        0].text = 'Теоретические и практические знания\n' + 'РО 1. Определять основные понятия и отрывки для планирования, разработки и проведения исследований\n' + 'РО 2. Описывать фундаментальные этические стандарты научных исследований\n''РО 3. Определять различные методологии\n' + 'РО 4. Объяснять и применять различные методы исследования (качественные и количественные)\n' + 'РО 5. Планировать и структурировать исследования в своей области исследований\n'
    table.rows[1].cells[
        1].text = 'ON5 описывать и применять основные методы и инструменты научного исследования, владение математическими и экономико-статистическими и финансовыми методами и инструментами для подготовки и проведения финансового анализа и оценки эффективности операционной, финансовой и инвестиционной деятельности компании;'
    table.rows[2].cells[
        0].text = 'Когнитивные и практические навыки и компетенции\n' + 'РО 6. Разработка планов исследований, работающих индивидуально и / или в группе\n' + 'РО 7. Применение различных методов сбора данных (библиографические исследования, исследования в режиме онлайн, интервью, опросы)\n' + 'РО 8. Применение различных качественных и количественных подходов к анализу данных;\n' + 'РО 9. Организация в эффективном представлении информации (защита планов исследования)\n'
    table.rows[2].cells[
        1].text = 'ON10 применять базовые исследовательские навыки, информационные и финансовые технологии, навыки критического мышления, коммуникационные навыки для выбора подходящих теорий и методологий, получения актуальной и точной информации, анализа данных и разработки выводов в теоретических исследованиях или прикладных проектах.'

    document.add_paragraph('\n\n3.	Тематический план')
    table = document.add_table(rows=len(modules)+1, cols=7)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Неделя'
    table.rows[0].cells[1].text = 'Тема / модуль'
    table.rows[0].cells[2].text = 'РО курса, РО ОП'
    table.rows[0].cells[3].text = 'Вопросы по теме / модулю'
    table.rows[0].cells[4].text = 'Задания'
    table.rows[0].cells[5].text = 'Литература'
    table.rows[0].cells[6].text = 'Структура оценок'
    j = 1
    for i in modules:
        table.rows[j].cells[0].text = str(i.week)
        table.rows[j].cells[1].text = i.theme
        table.rows[j].cells[2].text = 'РО курса, РО ОП'
        table.rows[j].cells[3].text = i.questions
        table.rows[j].cells[4].text = i.tasks
        table.rows[j].cells[5].text = i.literature.literature.title
        table.rows[j].cells[6].text = i.grading
        j+=1

    document.add_paragraph('\n\n4.	Система оценивания курса')
    able = document.add_table(rows=len(modules)+1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Тема / модуль'
    table.rows[0].cells[1].text = 'Максимальный процент(%)'
    table.rows[0].cells[2].text = 'Максимальный вес(%)'
    table.rows[0].cells[3].text = 'Итого в баллах'
    j=1
    for i in modules:
        table.rows[j].cells[0].text = i.theme
        table.rows[j].cells[1].text = str(i.max_percent)
        table.rows[j].cells[2].text = str(i.max_weight)
        table.rows[j].cells[3].text = str(i.total_in_points)
        j+=1

    document.add_paragraph('\n\n5.	Список литературы')

    document.add_paragraph("""Обязательная литература
        1.	Новиков, А.М. Методология научного исследования [Текст]: учебно-методическое пособие/ А.М. Новиков, Д.А. Новиков. – Изд. 2-е, - Москва: Кн.дом “ЛИБРОКОМ”:URSS, 2013. -270 с. 
        2.	An introduction to Business research methods Dr. Sue Greener, Dr. Joe Martell.- 2nd. ed.- Bookboon. com., 2015.- 137 p. 
        3.	Герасимов Б.И. Основы научных исследований [Электронный ресурс] / Б.И. Герасимов, В.В. Дробышева, Н.В. Злобина и др. - М.: Форум: НИЦ Инфра-М, 2013. Режим доступа: http://znanium.com/bookread.php?book=390595 (дата обращения 02.09.2016)

    Дополнительная литература 
    4.	Robert, K.Yin.Case Study Research [Текст]: Design and Methods / K.Yin Robert.- USA: Sage, 2014.- 282 с.
    5.	Орехов А.М. Методы экономических исследований [Электронный ресурс]: Учебное пособие / А.М. Орехов. - 2-e изд. - М.: НИЦ Инфра-М, 2013. - 344 с. Режим доступа: http://znanium.com/bookread.php?book=362627 (дата обращения 15.03.2015) 
    6.	Мильчакова, Н.Н., Яркова, Е.Н. Методы социально-экономических исследований: учебное пособие/ Н.Н. Мильчакова, Е. Н. Яркова; Тюм. гос. ун-т. - Тюмень: Изд-во ТюмГУ, 2014. - 379 с.

    Интернет ресурсы
        1.	Библиотека AlmaU  http://lib.almau.edu.kz/
        2.	Научная электронная библиотека http://elibrary.ru/ 
        3.	Научно-образовательный портал: http://www.med-edu.ru/
        4.	Международные организации Организация Объединенных Наций(ООН)- United Nations(UN)- http://www.un.org/ 
        5.	Международный валютный фонд(МВФ)- International Monetary Fund – IMFhttp://www.imf.org Всемирный Банк (World Bank)- http://www.worldbank.org 
        6.	Всемирная организация интеллектуальной собственности (ВОИС)- World Intellectual Property Organization (WIPO) - http://www.wipo.org 
        7.	Всемирный экономический форум - World Economic Forum- http://www.weforum.org БРИКС http://infobrics.org ШОС http://infoshos.ru 
        8.	Национальный банк РК. Официальный интернет ресурс http://www.nationalbank.kz/?switch=russian
        9.	Казахстанская фондовая биржа (KASE) Официальный интернет ресурс kase.kz
        10.	Кафедра экономической методологии и истории Высшей школы экономики: курсы, публикации http://www.hse.ru/kafedry/economy/ec_methodology_history/default.htm 
        11.	Портал по социологии, экономике и менеджменту www.ecsocman.edu.ru 
        12.	Портал по общественным наукам www.socionet.ru 

    """)

    document.add_paragraph("""6.	Философия преподавания и обучения
        Процесс обучения основывается на освоении теоретического материала на лекциях, на самостоятельном изучении материалов, практического применения знаний и обсуждениях в аудитории. Студенты, обучаясь в условиях использования активных форм, работая в группах, решая конкретные ситуационные задачи, приобретут способность при¬нимать решения в нестандартных ситуациях, умение работать в команде, самостоятельно добывать, анализировать и эффективно использовать информацию, рационально работать.
        Задача преподавателя будет заключаться в том, чтобы обеспечить учебным материалом, рекомендуемой литературой, донести сложные аспекты в доступной форме. Преподаватель несет ответственность за успешное освоение знаний и навыков в течение контактных часов и в процессе руководства самостоятельной работой студентов
        Подведение итогов преподавателем в конце недели позволяет студентам видеть свои еженедельные результаты, образующие средневзвешенные оценки уровня достижений (GPI).
    """)

    document.add_paragraph("""7.	Политика курса
    Этика занятий
    Освоение дисциплины «Research Methods» предусматривает 
    -	обязательное посещение занятий;
    -	активность во время занятий;
    -	подготовка к занятиям, выполнение домашнего задания;
    -	сдача заданий в установленные сроки;
    -	быть терпимым, открытым и доброжелательным; 
    -	конструктивно поддерживать обратную связь на всех занятиях; 
    -	быть пунктуальным и обязательным.
    Недопустимо:
    -	пропуски по неуважительным причинам;
    -	опоздание и уход с занятий (В случае опоздания студент не допускается на занятие, т.к. он нарушает ход учебного занятия);
    -	несвоевременная сдача заданий и др. 
    При пропусках занятий по уважительной причине допускается отработка пройденного материала.

    Этика экзамена
    Недопустимо:
    -	опоздание; 
    -	пользование мобильными телефонами во время экзамена; 
    -	списывание при сдаче экзамена. За списывание на контрольном мероприятии студент удаляется из аудитории и ему выставляется 0 баллов.
    Если в силу каких-либо уважительных причин вы отсутствовали во время проведения контрольного мероприятия, вам предоставляется возможность пройти его в дополнительно назначенное преподавателем время (РК и ИК сдаются с разрешения декана), в противном случае вы получаете «0» баллов.
    Политика академического поведения и этики основана на Кодексе корпоративной культуры, Этическом кодексе студента, Правилах внутреннего распорядка AlmaU.

    Информация и связь
    Вы должны регулярно (ежедневно) проверять Личную страницу в автоматизированной информационной системе, LMS и электронную почту, чтобы получать дополнительную информацию, задания или знать изменения в расписании. 

    """)
    filepath = f"D:\\Загрузки\\{syllabus.syllabus_name}.docx"
    document.save(f"{syllabus.syllabus_name}.docx")
    try:
        pythoncom.CoInitializeEx(0)
        word = client.DispatchEx("Word.Application")
        target_path = filepath.replace(".docx", r".pdf")
        word_doc = word.Documents.Open(filepath)
        word_doc.SaveAs(target_path, FileFormat=17)
        word_doc.Close()
    except Exception as e:
        raise e
    finally:
        word.Quit()


    with open(target_path, 'rb') as pdf_file:
            pdf_data = pdf_file.read()

    # Send the PDF as response with appropriate content type to open in a new window
    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{syllabus.syllabus_name}.pdf"'
    return response

def download_syllabus_as_word_rus(request, syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    document = Document()
    document.sections[0].page_width = Inches(10)
    document.sections[0].left_margin = Inches(1 / 3)
    document.sections[0].right_margin = Inches(1 / 3)
    table = document.add_table(rows=1, cols=3)
    row = table.rows[0]
    cell1 = row.cells[0]
    image_path = "C:\\Users\\User\\Desktop\\gorizontal'niy rgb.png"
    cell1_paragraph = cell1.paragraphs[0]
    run = cell1_paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.0))
    cell2 = row.cells[2]
    text_content = "УТВЕРЖДАЮ\nДекан Школы/Центра\nФИO\n_________________\n«___» ______202____"
    cell2.text = text_content

    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    row = table.rows[0]
    row.cells[0].text = "КОД И НАЗВАНИЕ ДИСЦИПЛИНЫ:\n" + syllabus.syllabus_name + " " + syllabus.course.__str__()
    hours = syllabus.total_hours-syllabus.classroom_hours
    row.cells[1].text = "КРЕДИТЫ ECTS И ЧАСЫ:\n3 ECTS\nВсего часов:" + str(syllabus.total_hours) + "\nАудиторные часы:" + str(syllabus.classroom_hours) + "часов\nСамостоятельная работа\n(СРОП, СРО):" +  str(hours) + "часов"

    cell13 = row.cells[2]
    text_content = "ПРЕРЕКВИЗИТЫ:\n" + syllabus.prerequisites
    cell13.text = text_content

    row = table.rows[1]
    cell11 = row.cells[0]
    text_content = "УРОВЕНЬ ОБУЧЕНИЯ:\n" + str(syllabus.training_level)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "СЕМЕСТР:\n" + str(syllabus.semester)
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ОБРАЗОВАТЕЛЬНАЯ ПРОГРАММА: \n" + str(syllabus.edu_programms)
    cell13.text = text_content

    row = table.rows[2]
    cell11 = row.cells[0]
    text_content = "ЯЗЫК ОБУЧЕНИЯ:\n" + str(syllabus.language_of_education)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "УРОВЕНЬ ВЛАДЕНИЯ ЯЗЫКОМ ОБУЧЕНИЯ:\n" + str(syllabus.proficiency_level)
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ФОРМАТ ОБУЧЕНИЯ:\n" + str(syllabus.format_of_training)
    cell13.text = text_content

    row = table.rows[3]
    cell11 = row.cells[0]
    text_content = "ПРЕПОДАВАТЕЛЬ:\n" + str(syllabus.instructor)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "КОНТАКТЫ ПРЕПОДАВАТЕЛЯ:\nЭл.почта/ телефон:" + syllabus.instructor.email
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ВРЕМЯ И МЕСТО ПРОВЕДЕНИЯ ЗАНЯТИЙ:\n" + syllabus.time_place
    cell13.text = text_content

    text_to_write = "Цель курса\nДанная дисциплина нацелена на поддержание и развитие аналитического, критического мышления и творческих навыков, а также написания и презентации исследования."
    paragraph = document.add_paragraph(text_to_write)

    text_to_write = "График занятий и задания"
    paragraph = document.add_paragraph(text_to_write)
    modules = Module.objects.filter(syllabus=syllabus)
    table1 = document.add_table(rows=len(modules)+1, cols=4)
    table1.style = 'Table Grid'
    column = table1.columns[0]
    column.width = Inches(0.4)
    table1.rows[0].cells[0].text = 'Недели'
    table1.rows[0].cells[1].text = 'Тема / модуль'
    table1.rows[0].cells[2].text = 'Формат проведения занятий'
    table1.rows[0].cells[3].text = 'Задания'
    j = 1
    for i in modules:
        table1.rows[j].cells[0].text = str(i.week)
        table1.rows[j].cells[0].width = Inches(0.4)
        table1.rows[j].cells[1].text = i.theme
        table1.rows[j].cells[2].text = str(i.format.__str__())
        table1.rows[j].cells[3].text = i.tasks
        j+=1
    document.add_paragraph('\n\n\n')
    new_table1 = document.add_table(rows=1, cols=2)

    row = new_table1.rows[0]
    table1 = row.cells[0].add_table(rows=4, cols=1)
    table1.rows[0].cells[0].text = 'Academic Handbook'
    run = table1.rows[1].cells[0].paragraphs[0].add_run()
    table1.rows[2].cells[
        0].text = 'https://almauedu-my.sharepoint.com/:f:/g/personal/f_abdoldina_almau_edu_kz/EnVy7hCS47hMoVtpgjfq3-YBY2biThYahFoceoI9xY1n3A?e=wASl1u'
    table1.rows[3].cells[0].text = 'Составлено:\nк.э.н., lecturer	___________	ФИО ППС'
    table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = row.cells[1].add_table(rows=4, cols=1)
    table1.rows[0].cells[0].text = 'Результаты обучения курса. Задания и политика курса:'
    run = table1.rows[1].cells[0].paragraphs[0].add_run()
    table1.rows[2].cells[
        0].text = 'https://docs.google.com/document/u/0/d/19QyuM6a1uyAXd49Rb9cpOW43lyYe5Cc0/mobilebasic'
    table1.rows[3].cells[0].text = 'Согласовано:\nДиректор УМ	___________	Абдолдина Ф.Н.'
    table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('\n\n\n')
    document.add_paragraph('1.	Описание курса')
    document.add_paragraph(
        '	Дисциплина «Research Methods» предлагает студентам всестороннее представление о том, как проводить исследовательскую работу/проект и профессионально презентовать достигнутые результаты, учитывая специфику предстоящих научных исследований по специальностям, и формирует компетенции в области научных исследований. Переходя от первых шагов исследовательской работы/проекта (определение проблемных вопросов исследования) к последнему результату (разработка рекомендаций), в конце курса студенты смогут провести углубленное исследование в своей области, и смогут более уверенно представлять и защищать свои идеи перед критической аудиторией.' +
        '\n	Основная цель курса - формирование креативного исследовательского мышления и способностей решать разнообразные хозяйственные, социальные, психологические задачи путем использования современных методов, приемов и средств научного исследования.' +
        '\n	Теоретический подход сочетается с практическими заданиями по проводимым исследованиям и выбранной темы дипломного проекта.')
    document.add_paragraph('\n\n')
    document.add_paragraph(
        '2.	Таблица соответствия Результатов обучения курса Результатам обучения образовательной программы')

    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Результаты обучения курса (РО курса)'
    table.rows[0].cells[1].text = 'Результаты обучения образовательной программы (РО ОП)'
    table.rows[1].cells[
        0].text = 'Теоретические и практические знания\n' + 'РО 1. Определять основные понятия и отрывки для планирования, разработки и проведения исследований\n' + 'РО 2. Описывать фундаментальные этические стандарты научных исследований\n''РО 3. Определять различные методологии\n' + 'РО 4. Объяснять и применять различные методы исследования (качественные и количественные)\n' + 'РО 5. Планировать и структурировать исследования в своей области исследований\n'
    table.rows[1].cells[
        1].text = 'ON5 описывать и применять основные методы и инструменты научного исследования, владение математическими и экономико-статистическими и финансовыми методами и инструментами для подготовки и проведения финансового анализа и оценки эффективности операционной, финансовой и инвестиционной деятельности компании;'
    table.rows[2].cells[
        0].text = 'Когнитивные и практические навыки и компетенции\n' + 'РО 6. Разработка планов исследований, работающих индивидуально и / или в группе\n' + 'РО 7. Применение различных методов сбора данных (библиографические исследования, исследования в режиме онлайн, интервью, опросы)\n' + 'РО 8. Применение различных качественных и количественных подходов к анализу данных;\n' + 'РО 9. Организация в эффективном представлении информации (защита планов исследования)\n'
    table.rows[2].cells[
        1].text = 'ON10 применять базовые исследовательские навыки, информационные и финансовые технологии, навыки критического мышления, коммуникационные навыки для выбора подходящих теорий и методологий, получения актуальной и точной информации, анализа данных и разработки выводов в теоретических исследованиях или прикладных проектах.'

    document.add_paragraph('\n\n3.	Тематический план')
    table = document.add_table(rows=len(modules)+1, cols=7)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Неделя'
    table.rows[0].cells[1].text = 'Тема / модуль'
    table.rows[0].cells[2].text = 'РО курса, РО ОП'
    table.rows[0].cells[3].text = 'Вопросы по теме / модулю'
    table.rows[0].cells[4].text = 'Задания'
    table.rows[0].cells[5].text = 'Литература'
    table.rows[0].cells[6].text = 'Структура оценок'
    j = 1
    for i in modules:
        table.rows[j].cells[0].text = str(i.week)
        table.rows[j].cells[1].text = i.theme
        table.rows[j].cells[2].text = 'РО курса, РО ОП'
        table.rows[j].cells[3].text = i.questions
        table.rows[j].cells[4].text = i.tasks
        table.rows[j].cells[5].text = i.literature.literature.title
        table.rows[j].cells[6].text = i.grading
        j+=1

    document.add_paragraph('\n\n4.	Система оценивания курса')
    able = document.add_table(rows=len(modules)+1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Тема / модуль'
    table.rows[0].cells[1].text = 'Максимальный процент(%)'
    table.rows[0].cells[2].text = 'Максимальный вес(%)'
    table.rows[0].cells[3].text = 'Итого в баллах'
    j=1
    for i in modules:
        table.rows[j].cells[0].text = i.theme
        table.rows[j].cells[1].text = str(i.max_percent)
        table.rows[j].cells[2].text = str(i.max_weight)
        table.rows[j].cells[3].text = str(i.total_in_points)
        j+=1

    document.add_paragraph('\n\n5.	Список литературы')

    document.add_paragraph("""Обязательная литература
        1.	Новиков, А.М. Методология научного исследования [Текст]: учебно-методическое пособие/ А.М. Новиков, Д.А. Новиков. – Изд. 2-е, - Москва: Кн.дом “ЛИБРОКОМ”:URSS, 2013. -270 с. 
        2.	An introduction to Business research methods Dr. Sue Greener, Dr. Joe Martell.- 2nd. ed.- Bookboon. com., 2015.- 137 p. 
        3.	Герасимов Б.И. Основы научных исследований [Электронный ресурс] / Б.И. Герасимов, В.В. Дробышева, Н.В. Злобина и др. - М.: Форум: НИЦ Инфра-М, 2013. Режим доступа: http://znanium.com/bookread.php?book=390595 (дата обращения 02.09.2016)

    Дополнительная литература 
    4.	Robert, K.Yin.Case Study Research [Текст]: Design and Methods / K.Yin Robert.- USA: Sage, 2014.- 282 с.
    5.	Орехов А.М. Методы экономических исследований [Электронный ресурс]: Учебное пособие / А.М. Орехов. - 2-e изд. - М.: НИЦ Инфра-М, 2013. - 344 с. Режим доступа: http://znanium.com/bookread.php?book=362627 (дата обращения 15.03.2015) 
    6.	Мильчакова, Н.Н., Яркова, Е.Н. Методы социально-экономических исследований: учебное пособие/ Н.Н. Мильчакова, Е. Н. Яркова; Тюм. гос. ун-т. - Тюмень: Изд-во ТюмГУ, 2014. - 379 с.

    Интернет ресурсы
        1.	Библиотека AlmaU  http://lib.almau.edu.kz/
        2.	Научная электронная библиотека http://elibrary.ru/ 
        3.	Научно-образовательный портал: http://www.med-edu.ru/
        4.	Международные организации Организация Объединенных Наций(ООН)- United Nations(UN)- http://www.un.org/ 
        5.	Международный валютный фонд(МВФ)- International Monetary Fund – IMFhttp://www.imf.org Всемирный Банк (World Bank)- http://www.worldbank.org 
        6.	Всемирная организация интеллектуальной собственности (ВОИС)- World Intellectual Property Organization (WIPO) - http://www.wipo.org 
        7.	Всемирный экономический форум - World Economic Forum- http://www.weforum.org БРИКС http://infobrics.org ШОС http://infoshos.ru 
        8.	Национальный банк РК. Официальный интернет ресурс http://www.nationalbank.kz/?switch=russian
        9.	Казахстанская фондовая биржа (KASE) Официальный интернет ресурс kase.kz
        10.	Кафедра экономической методологии и истории Высшей школы экономики: курсы, публикации http://www.hse.ru/kafedry/economy/ec_methodology_history/default.htm 
        11.	Портал по социологии, экономике и менеджменту www.ecsocman.edu.ru 
        12.	Портал по общественным наукам www.socionet.ru 

    """)

    document.add_paragraph("""6.	Философия преподавания и обучения
        Процесс обучения основывается на освоении теоретического материала на лекциях, на самостоятельном изучении материалов, практического применения знаний и обсуждениях в аудитории. Студенты, обучаясь в условиях использования активных форм, работая в группах, решая конкретные ситуационные задачи, приобретут способность при¬нимать решения в нестандартных ситуациях, умение работать в команде, самостоятельно добывать, анализировать и эффективно использовать информацию, рационально работать.
        Задача преподавателя будет заключаться в том, чтобы обеспечить учебным материалом, рекомендуемой литературой, донести сложные аспекты в доступной форме. Преподаватель несет ответственность за успешное освоение знаний и навыков в течение контактных часов и в процессе руководства самостоятельной работой студентов
        Подведение итогов преподавателем в конце недели позволяет студентам видеть свои еженедельные результаты, образующие средневзвешенные оценки уровня достижений (GPI).
    """)

    document.add_paragraph("""7.	Политика курса
    Этика занятий
    Освоение дисциплины «Research Methods» предусматривает 
    -	обязательное посещение занятий;
    -	активность во время занятий;
    -	подготовка к занятиям, выполнение домашнего задания;
    -	сдача заданий в установленные сроки;
    -	быть терпимым, открытым и доброжелательным; 
    -	конструктивно поддерживать обратную связь на всех занятиях; 
    -	быть пунктуальным и обязательным.
    Недопустимо:
    -	пропуски по неуважительным причинам;
    -	опоздание и уход с занятий (В случае опоздания студент не допускается на занятие, т.к. он нарушает ход учебного занятия);
    -	несвоевременная сдача заданий и др. 
    При пропусках занятий по уважительной причине допускается отработка пройденного материала.

    Этика экзамена
    Недопустимо:
    -	опоздание; 
    -	пользование мобильными телефонами во время экзамена; 
    -	списывание при сдаче экзамена. За списывание на контрольном мероприятии студент удаляется из аудитории и ему выставляется 0 баллов.
    Если в силу каких-либо уважительных причин вы отсутствовали во время проведения контрольного мероприятия, вам предоставляется возможность пройти его в дополнительно назначенное преподавателем время (РК и ИК сдаются с разрешения декана), в противном случае вы получаете «0» баллов.
    Политика академического поведения и этики основана на Кодексе корпоративной культуры, Этическом кодексе студента, Правилах внутреннего распорядка AlmaU.

    Информация и связь
    Вы должны регулярно (ежедневно) проверять Личную страницу в автоматизированной информационной системе, LMS и электронную почту, чтобы получать дополнительную информацию, задания или знать изменения в расписании. 

    """)
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"{syllabus.syllabus_name}.docx"
    response = FileResponse(file_stream, as_attachment=True, filename=filename)
    return response


def download_syllabus_as_word_kz(request, syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    document = Document()
    document.sections[0].page_width = Inches(10)
    document.sections[0].left_margin = Inches(1 / 3)
    document.sections[0].right_margin = Inches(1 / 3)
    table = document.add_table(rows=1, cols=3)
    row = table.rows[0]
    cell1 = row.cells[0]
    image_path = "C:\\Users\\User\\Desktop\\gorizontal'niy rgb.png"
    cell1_paragraph = cell1.paragraphs[0]
    run = cell1_paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.0))
    cell2 = row.cells[2]
    text_content = "БЕКІТЕМІН\nМектеп/Орталық деканы \nТАӘ\n_________________\n«___» ______202____"
    cell2.text = text_content

    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    row = table.rows[0]
    row.cells[0].text = "ПӘННІҢ АТЫ МЕН КОДЫ:\n" + syllabus.syllabus_name + " " + syllabus.course.__str__()
    hours = syllabus.total_hours-syllabus.classroom_hours
    row.cells[1].text = "ECTS КРЕДИТТЕРІ ЖӘНЕ САҒАТТАР:\n3 ECTS\nСағаттар барлығы:" + str(syllabus.total_hours) + "\nАудиторлық сағаттар:" + str(syllabus.classroom_hours) + "сағат\nӨзіндік жұмыс (ООӨЖ, ОӨЖ):" +  str(hours) + "сағат"

    cell13 = row.cells[2]
    text_content = "ПРЕРЕКВИЗИТТЕР:\n" + syllabus.prerequisites
    cell13.text = text_content

    row = table.rows[1]
    cell11 = row.cells[0]
    text_content = "ОҚУ ДЕҢГЕЙІ:\n" + str(syllabus.training_level)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "СЕМЕСТР:\n" + str(syllabus.semester)
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "БІЛІМ БЕРУ БАҒДАРЛАМАСЫ: \n" + str(syllabus.edu_programms)
    cell13.text = text_content

    row = table.rows[2]
    cell11 = row.cells[0]
    text_content = "ОҚУ ТІЛІ:\n" + str(syllabus.language_of_education)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "ОҚУ ТІЛІН МЕҢГЕРУ ДЕҢГЕЙІ:\n" + str(syllabus.proficiency_level)
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "ОҚУ ФОРМАТЫ:\n" + str(syllabus.format_of_training)
    cell13.text = text_content

    row = table.rows[3]
    cell11 = row.cells[0]
    text_content = "ОҚЫТУШЫ:\n" + str(syllabus.instructor)
    cell11.text = text_content

    cell12 = row.cells[1]
    text_content = "ОҚЫТУШЫМЕН БАЙЛАНЫС:\nЭл.пошта:" + syllabus.instructor.email
    cell12.text = text_content

    cell13 = row.cells[2]
    text_content = "САБАҚТЫҢ УАҚЫТЫ МЕН ОРНЫ:\n" + syllabus.time_place
    cell13.text = text_content

    text_to_write = "Курстың мақсаты\n......"
    paragraph = document.add_paragraph(text_to_write)

    text_to_write = "Сабақ кестесі және тапсырмалар"
    paragraph = document.add_paragraph(text_to_write)
    modules = Module.objects.filter(syllabus=syllabus)
    table1 = document.add_table(rows=len(modules)+1, cols=4)
    table1.style = 'Table Grid'
    column = table1.columns[0]
    column.width = Inches(0.4)
    table1.rows[0].cells[0].text = 'Апталар'
    table1.rows[0].cells[1].text = 'Тақырып / модуль'
    table1.rows[0].cells[2].text = 'Сабақты жүргізу форматы'
    table1.rows[0].cells[3].text = 'Тапсырмалар'
    j = 1
    for i in modules:
        table1.rows[j].cells[0].text = str(i.week)
        table1.rows[j].cells[0].width = Inches(0.4)
        table1.rows[j].cells[1].text = i.theme
        table1.rows[j].cells[2].text = str(i.format.__str__())
        table1.rows[j].cells[3].text = i.tasks
        j+=1
    document.add_paragraph('\n\n\n')
    new_table1 = document.add_table(rows=1, cols=2)

    row = new_table1.rows[0]
    table1 = row.cells[0].add_table(rows=4, cols=1)
    table1.rows[0].cells[0].text = 'Academic Handbook'
    run = table1.rows[1].cells[0].paragraphs[0].add_run()
    table1.rows[2].cells[
        0].text = 'https://almauedu-my.sharepoint.com/:f:/g/personal/f_abdoldina_almau_edu_kz/EnVy7hCS47hMoVtpgjfq3-YBY2biThYahFoceoI9xY1n3A?e=wASl1u'
    table1.rows[3].cells[0].text = 'Құрастырған:\nк.э.н., lecturer	___________	ПОҚ ТАӘ'
    table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = row.cells[1].add_table(rows=4, cols=1)
    table1.rows[0].cells[0].text = 'Курсты оқыту нәтижелері. Курс тапсырмалары мен саясаты:'
    run = table1.rows[1].cells[0].paragraphs[0].add_run()
    table1.rows[2].cells[
        0].text = 'https://docs.google.com/document/u/0/d/19QyuM6a1uyAXd49Rb9cpOW43lyYe5Cc0/mobilebasic'
    table1.rows[3].cells[0].text = 'Келісілген:\nӘБ директоры	___________	Абдолдина Ф.Н.'
    table1.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1.rows[2].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('\n\n\n')
    document.add_paragraph('1.	Курс сипаттамасы\n..........')
    # document.add_paragraph(
    #     '	Дисциплина «Research Methods» предлагает студентам всестороннее представление о том, как проводить исследовательскую работу/проект и профессионально презентовать достигнутые результаты, учитывая специфику предстоящих научных исследований по специальностям, и формирует компетенции в области научных исследований. Переходя от первых шагов исследовательской работы/проекта (определение проблемных вопросов исследования) к последнему результату (разработка рекомендаций), в конце курса студенты смогут провести углубленное исследование в своей области, и смогут более уверенно представлять и защищать свои идеи перед критической аудиторией.' +
    #     '\n	Основная цель курса - формирование креативного исследовательского мышления и способностей решать разнообразные хозяйственные, социальные, психологические задачи путем использования современных методов, приемов и средств научного исследования.' +
    #     '\n	Теоретический подход сочетается с практическими заданиями по проводимым исследованиям и выбранной темы дипломного проекта.')
    document.add_paragraph('\n\n')
    document.add_paragraph(
        '2.	Курс оқыту Нәтижелерінің білім беру бағдарламасын оқыту Нәтижелеріне сәйкестік кестесі')

    table = document.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Курсты оқыту нәтижелері (Курсты ОН)'
    table.rows[0].cells[1].text = 'Білім беру бағдарламасын оқыту нәтижелері (БББ ОН)'
    table.rows[1].cells[
        0].text = 'Теоретические и практические знания\n' + 'РО 1. Определять основные понятия и отрывки для планирования, разработки и проведения исследований\n' + 'РО 2. Описывать фундаментальные этические стандарты научных исследований\n''РО 3. Определять различные методологии\n' + 'РО 4. Объяснять и применять различные методы исследования (качественные и количественные)\n' + 'РО 5. Планировать и структурировать исследования в своей области исследований\n'
    table.rows[1].cells[
        1].text = 'ON5 описывать и применять основные методы и инструменты научного исследования, владение математическими и экономико-статистическими и финансовыми методами и инструментами для подготовки и проведения финансового анализа и оценки эффективности операционной, финансовой и инвестиционной деятельности компании;'
    table.rows[2].cells[
        0].text = 'Когнитивные и практические навыки и компетенции\n' + 'РО 6. Разработка планов исследований, работающих индивидуально и / или в группе\n' + 'РО 7. Применение различных методов сбора данных (библиографические исследования, исследования в режиме онлайн, интервью, опросы)\n' + 'РО 8. Применение различных качественных и количественных подходов к анализу данных;\n' + 'РО 9. Организация в эффективном представлении информации (защита планов исследования)\n'
    table.rows[2].cells[
        1].text = 'ON10 применять базовые исследовательские навыки, информационные и финансовые технологии, навыки критического мышления, коммуникационные навыки для выбора подходящих теорий и методологий, получения актуальной и точной информации, анализа данных и разработки выводов в теоретических исследованиях или прикладных проектах.'

    document.add_paragraph('\n\n3.	Тақырыптық жоспар')
    table = document.add_table(rows=len(modules)+1, cols=7)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Апта'
    table.rows[0].cells[1].text = 'Тақырып/ модуль'
    table.rows[0].cells[2].text = 'Курсты ОН, БББ ОН'
    table.rows[0].cells[3].text = 'Тақырып / модуль бойынша сұрақтар'
    table.rows[0].cells[4].text = 'Тапсырмалар'
    table.rows[0].cells[5].text = 'Әдебиет'
    table.rows[0].cells[6].text = 'Бағалау құрылымы'
    j = 1
    for i in modules:
        table.rows[j].cells[0].text = str(i.week)
        table.rows[j].cells[1].text = i.theme
        table.rows[j].cells[2].text = 'РО курса, РО ОП'
        table.rows[j].cells[3].text = i.questions
        table.rows[j].cells[4].text = i.tasks
        table.rows[j].cells[5].text = i.literature.literature.title
        table.rows[j].cells[6].text = i.grading
        j+=1

    document.add_paragraph('\n\n4.	Курсты бағалау жүйесі')
    able = document.add_table(rows=len(modules)+1, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Тақырып / модуль'
    table.rows[0].cells[1].text = 'Максималды пайыз(%)'
    table.rows[0].cells[2].text = 'Максималды салмақ(%)'
    table.rows[0].cells[3].text = 'Барлығы баллмен'
    j=1
    for i in modules:
        table.rows[j].cells[0].text = i.theme
        table.rows[j].cells[1].text = str(i.max_percent)
        table.rows[j].cells[2].text = str(i.max_weight)
        table.rows[j].cells[3].text = str(i.total_in_points)
        j+=1

    document.add_paragraph('\n\n5.	Әдебиеттер тізімі')

    document.add_paragraph("""Қажетті әдебиет
        1.	Новиков, А.М. Методология научного исследования [Текст]: учебно-методическое пособие/ А.М. Новиков, Д.А. Новиков. – Изд. 2-е, - Москва: Кн.дом “ЛИБРОКОМ”:URSS, 2013. -270 с. 
        2.	An introduction to Business research methods Dr. Sue Greener, Dr. Joe Martell.- 2nd. ed.- Bookboon. com., 2015.- 137 p. 
        3.	Герасимов Б.И. Основы научных исследований [Электронный ресурс] / Б.И. Герасимов, В.В. Дробышева, Н.В. Злобина и др. - М.: Форум: НИЦ Инфра-М, 2013. Режим доступа: http://znanium.com/bookread.php?book=390595 (дата обращения 02.09.2016)

    Қосымша әдебиет 
    4.	Robert, K.Yin.Case Study Research [Текст]: Design and Methods / K.Yin Robert.- USA: Sage, 2014.- 282 с.
    5.	Орехов А.М. Методы экономических исследований [Электронный ресурс]: Учебное пособие / А.М. Орехов. - 2-e изд. - М.: НИЦ Инфра-М, 2013. - 344 с. Режим доступа: http://znanium.com/bookread.php?book=362627 (дата обращения 15.03.2015) 
    6.	Мильчакова, Н.Н., Яркова, Е.Н. Методы социально-экономических исследований: учебное пособие/ Н.Н. Мильчакова, Е. Н. Яркова; Тюм. гос. ун-т. - Тюмень: Изд-во ТюмГУ, 2014. - 379 с.

    Интернет ресурстар
        1.	Библиотека AlmaU  http://lib.almau.edu.kz/
        2.	Научная электронная библиотека http://elibrary.ru/ 
        3.	Научно-образовательный портал: http://www.med-edu.ru/
        4.	Международные организации Организация Объединенных Наций(ООН)- United Nations(UN)- http://www.un.org/ 
        5.	Международный валютный фонд(МВФ)- International Monetary Fund – IMFhttp://www.imf.org Всемирный Банк (World Bank)- http://www.worldbank.org 
        6.	Всемирная организация интеллектуальной собственности (ВОИС)- World Intellectual Property Organization (WIPO) - http://www.wipo.org 
        7.	Всемирный экономический форум - World Economic Forum- http://www.weforum.org БРИКС http://infobrics.org ШОС http://infoshos.ru 
        8.	Национальный банк РК. Официальный интернет ресурс http://www.nationalbank.kz/?switch=russian
        9.	Казахстанская фондовая биржа (KASE) Официальный интернет ресурс kase.kz
        10.	Кафедра экономической методологии и истории Высшей школы экономики: курсы, публикации http://www.hse.ru/kafedry/economy/ec_methodology_history/default.htm 
        11.	Портал по социологии, экономике и менеджменту www.ecsocman.edu.ru 
        12.	Портал по общественным наукам www.socionet.ru 

    """)

    document.add_paragraph("""6.	Білім беру мен оқыту философиясы
        Процесс обучения основывается на освоении теоретического материала на лекциях, на самостоятельном изучении материалов, практического применения знаний и обсуждениях в аудитории. Студенты, обучаясь в условиях использования активных форм, работая в группах, решая конкретные ситуационные задачи, приобретут способность при¬нимать решения в нестандартных ситуациях, умение работать в команде, самостоятельно добывать, анализировать и эффективно использовать информацию, рационально работать.
        Задача преподавателя будет заключаться в том, чтобы обеспечить учебным материалом, рекомендуемой литературой, донести сложные аспекты в доступной форме. Преподаватель несет ответственность за успешное освоение знаний и навыков в течение контактных часов и в процессе руководства самостоятельной работой студентов
        Подведение итогов преподавателем в конце недели позволяет студентам видеть свои еженедельные результаты, образующие средневзвешенные оценки уровня достижений (GPI).
    """)

    document.add_paragraph("""7.	Курс саясаты
    Сабақ этикасы
    Емтихан этикасы
    Ақпарат пен байланыс
    """)

    document.save('demo.docx')

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"{syllabus.syllabus_name}.docx"
    response = FileResponse(file_stream, as_attachment=True, filename=filename)
    return response


def download_syllabus_as_pdf(request, syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer)

    p.setFont("Helvetica", 12)
    y = 650

    # Helper function to draw text with specified position and move the y-coordinate
    def draw_text(x, text):
        nonlocal y
        p.drawString(x, y, text)
        y -= 20

    draw_text(100, syllabus.syllabus_name)
    draw_text(100, f"Discipline: {syllabus.course}")
    draw_text(100, f"Training Level: {syllabus.training_level}")
    draw_text(100, f"Language of Education: {syllabus.language_of_education}")
    draw_text(100, f"Language Proficiency Level: {syllabus.proficiency_level}")
    draw_text(100, f"Total Hours: {syllabus.total_hours}")
    draw_text(100, f"Classroom Hours: {syllabus.classroom_hours}")
    draw_text(100, f"Semester: {syllabus.semester}")
    draw_text(100, f"ECTS Credits: {syllabus.ects}")
    draw_text(100, f"IW Hours: {syllabus.iw_hours}")
    draw_text(100, f"Prerequisites: {syllabus.prerequisites}")
    draw_text(100, f"Training Format: {syllabus.format_of_training}")
    draw_text(100, f"Educational Programs: {syllabus.edu_programms}")
    draw_text(100, f"Time and Place of Conduct: {syllabus.time_place}")
    draw_text(100, f"Instructor/Teacher: {syllabus.instructor}")
    draw_text(100, f"Course Objective: {syllabus.course_objective}")
    draw_text(100, f"Course Philosophy: {syllabus.course_philosophy}")
    draw_text(100, f"Course Policy: {syllabus.course_etics}")

    draw_text(100, "Literature")
    for literature in syllabus.literature_set.all():
        draw_text(120, literature.title)

    draw_text(100, "Modules")
    for module in syllabus.module_set.all():
        draw_text(120, f"Week {module.week}")
        draw_text(120, f"Theme: {module.theme}")
        draw_text(120, f"Format: {module.format}")
        draw_text(120, f"Tasks: {module.tasks}")
        draw_text(120, f"Course Learning Outcomes: {module.course_lo}")
        draw_text(120, f"Module Questions: {module.questions}")
        draw_text(120, f"Grading: {module.grading}")
        draw_text(120, f"Maximum Percentage: {module.max_percent}")
        draw_text(120, f"Maximum Weight: {module.max_weight}")
        draw_text(120, f"In Points: {module.total_in_points}")

    p.showPage()
    p.save()

    buffer.seek(0)

    response = FileResponse(buffer, as_attachment=True, filename=f"{syllabus.syllabus_name}.pdf")
    return response



# Create your views here.
def home(request):
    return render(request, 'syllabuses/home.html', {})



def login_v(request):
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('create_syllabus')  # Исправленный путь для перенаправления
            else:
                print("Why is this not returned for inval")
    else:
        form = AuthenticationForm()
    return render(request, 'syllabuses/login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('../create_syllabus')


























def create_syllabus(request):
    if request.method == 'POST':
        form = SyllabusForm(request.POST)
        if form.is_valid():
            syllabus = form.save(commit=False)
            syllabus.status = Status.objects.get(type="Created")
            syllabus.save()
            return redirect(f'literature_form/{syllabus.id}')
    else:
        form = SyllabusForm()
    return render(request, 'syllabuses/create_syllabus.html', {'form': form})



def next_step(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added literature")
        syllabus.save()
        l = request.POST["liter"]
        literature = Literature.objects.get(pk=l)
        mandatory = request.POST["mandatory"]
        LiteratureInSyllabus.objects.create(
            syllabus = syllabus,
            literature = literature,
            mandatory = mandatory
        )
    return render(request, 'syllabuses/next_step.html',{
                      'syllabus': syllabus, 
                      'literatures': Literature.objects.filter(course=syllabus.course),
                      'literaturesinsyllabus': LiteratureInSyllabus.objects.filter(syllabus = syllabus)
                    })

def delete_literature(request, pk, syllabus_id):
    literature = LiteratureInSyllabus.objects.get(pk=pk)
    syllabus_id=syllabus_id
    literature.delete()
    return redirect(f'../../../next_step/{syllabus_id}')


def delete_syllabus(request, syllabus_id):
    syl = Syllabus.objects.get(pk=syllabus_id)
    syl.delete()
    return redirect(f'../../../my_syllabuses')

def delete_module(request, pk, syllabus_id):
    module = Module.objects.get(pk=pk)
    syllabus_id=syllabus_id
    module.delete()
    return redirect(f'../../../add_module/{syllabus_id}')

def add_literature(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    # literature_form = SecondStepForm()
    if request.method == "POST":
        title = request.POST["title"]
        Literature.objects.create(
            course=syllabus.course,
            title=title,
        )


    return render(request, 'syllabuses/literature_form.html', 
                  {
                      'syllabus': syllabus, 
                      'literatures': Literature.objects.filter(course=syllabus.course),
                    })



def syllabus_details(request, syllabus_id: int):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)
    lo11 = CourseLO.objects.filter(syllabus=syllabus).filter(type=True)
    lo22 = CourseLO.objects.filter(syllabus=syllabus).filter(type=False)
    modules = Module.objects.filter(syllabus=syllabus).order_by('week')

    if request.method == 'POST':
        syllabus.syllabus_name= request.POST.get('syllabus_name')
        course = request.POST.get('course')
        training_level = request.POST.get('training_level')
        language_of_education = request.POST.get('language_of_education')
        proficiency_level = request.POST.get('proficiency_level')
        total_hours = request.POST.get('total_hours')
        classroom_hours = request.POST.get('classroom_hours')
        semester = request.POST.get('semester')
        ects = request.POST.get('ects')
        iw_hours = request.POST.get('iw_hours')
        prerequisites = request.POST.get('prerequisites')
        format_of_training = request.POST.get('format_of_training')
        edu_programms = request.POST.get('edu_programms')
        time_place = request.POST.get('time_place')
        instructor = request.POST.get('instructor')
        course_objective = request.POST.get('course_objective')
        agreed_with = request.POST.get('agreed_with')
        status = request.POST.get('status')
        course_philosophy = request.POST.get('course_philosophy')
        course_etics = request.POST.get('course_etics')
        asu = request.POST.get('asu')
        syllabus.course = Course.objects.get(pk=course)
        syllabus.training_level = EduLevel.objects.get(pk=training_level)
        syllabus.language_of_education = Language.objects.get(pk=language_of_education)
        syllabus.proficiency_level = Proficiency.objects.get(pk=proficiency_level)
        syllabus.total_hours = total_hours
        syllabus.classroom_hours = classroom_hours
        syllabus.semester = semester
        syllabus.ects = ects
        syllabus.iw_hours = iw_hours
        syllabus.prerequisites = prerequisites
        syllabus.format_of_training = Format.objects.get(pk=format_of_training)
        syllabus.edu_programms = edu_programms
        syllabus.time_place = time_place
        syllabus.instructor = CustomUser.objects.get(pk=instructor)
        syllabus.course_objective = course_objective
        syllabus.agreed_with = Director.objects.get(pk=agreed_with)
        syllabus.status = Status.objects.get(pk=status)
        syllabus.course_philosophy = course_philosophy
        syllabus.course_etics = course_etics
        if asu=="on":
            syllabus.asu = True
        else: 
            syllabus.asu = False
        syllabus.save()
        return redirect(reverse('syllabus_details', kwargs={'syllabus_id': syllabus_id}))

    
    
    return render(request, 'syllabuses/syllabus_details.html', {
        'syllabus': syllabus,
        'literatures': LiteratureInSyllabus.objects.filter(syllabus=syllabus),
        'lo11': lo11,
        'lo22': lo22,
        'modules': modules,
        'courses': Course.objects.all(),
        'edu_levels': EduLevel.objects.all(),
        'languages': Language.objects.all(),
        'proficiencies': Proficiency.objects.all(),
        'formats': Format.objects.all(),
        'instructors': CustomUser.objects.all(),
        'directors': Director.objects.all(),
        'statuses': Status.objects.all(),
    })





def half(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
  # literature_form = SecondStepForm()
    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added lo")
        syllabus.save()
        lo = request.POST["lo"]
        lo2 = request.POST["lo2"]
        CourseLO.objects.create(
            syllabus=syllabus,
            type = True,
            info = lo
        )
        CourseLO.objects.create(
            syllabus=syllabus,
            type = False,
            info = lo2
        )


    return render(request, 'syllabuses/half.html', 
                  {
                      'syllabus': syllabus, 

                    })

def edit_profile(request):
    currentuser = request.user

    return render(request, 'syllabuses/edit_profile.html', 
                  {
                      'user': currentuser, 

                    })




def add_policy(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added policy")
        syllabus.course_philosophy = request.POST["phylosophy"]
        syllabus.course_etics = request.POST["policy"]
        syllabus.save()
        messages.success(request, "Силлабус создан")
        return redirect('home')  # Перенаправление на главный экран

    return render(request, 'syllabuses/add_policy.html',
                   {
                       'syllabus': syllabus
                       })



def add_module(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)

    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added modules")
        syllabus.save()
        week = request.POST["week"]
        theme = request.POST["theme"]
        format = Format.objects.get(pk = request.POST["format"]) 
        tasks = request.POST["tasks"]
        lo = request.POST["lo"]
        questions = request.POST["questions"]
        liter = LiteratureInSyllabus.objects.get(pk = request.POST["liter"]) 
        grading = request.POST["grading"]
        maxpercent = request.POST["maxpercent"]
        maxvalue = request.POST["maxvalue"]
        total_in_points = request.POST["total_in_points"]

        Module.objects.create(
            syllabus=syllabus,
            week=week,
            tasks=tasks,
            course_lo=lo,
            theme=theme,
            format=format,
            questions=questions,
            literature=liter,
            grading=grading,
            max_percent=maxpercent,
            max_weight=maxvalue,
            total_in_points=total_in_points,
        )


    return render(request, 'syllabuses/add_module.html', 
                  {
                      'syllabus': syllabus, 
                      'literatures': LiteratureInSyllabus.objects.filter(syllabus=syllabus),
                      'formats': Format.objects.all(),
                      'modules': Module.objects.filter(syllabus=syllabus).order_by('week'),
                      'lo11': CourseLO.objects.get(syllabus=syllabus, type=True),
                      'lo22': CourseLO.objects.get(syllabus=syllabus, type=False),
                    })



def continue_edit(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    print(syllabus.status)

    if syllabus.status.type == "Created":
        redirect_url = f'../../literature_form/{syllabus_id}'
    elif syllabus.status.type == "added literature":
        redirect_url = f'../../half/{syllabus_id}'
    elif syllabus.status.type == "added lo":
        redirect_url = f'../../add_module/{syllabus_id}'
    else:
        redirect_url = reverse('syllabus_details', args=[syllabus_id])

    return render(request, 'continue_edit.html', {'redirect_url': redirect_url})





@staff_member_required
def my_syllabuses(request):
    syllabuses = Syllabus.objects.all()  # Получаем все силлабусы

    return render(request, 'syllabuses/my_syllabuses.html', {'syllabuses': syllabuses})



def view_syllabuses(request):
    user = request.user
    syllabuses = user.get_created_syllabuses()
    context = {
        'syllabuses': syllabuses
    }
    return render(request, 'syllabuses/view_syllabuses.html', context)

   

def add_instructor(request):
    if request.method == 'POST':
        form = CustomUserForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('../create_syllabus')  # Перенаправьте пользователя на страницу успешного добавления преподавателя
    else:
        form = CustomUserForm()
    return render(request, 'syllabuses/add_instructor.html', {'form': form})

class SchoolView(View):
    def get(self, request):
        schools = School.objects.all()
        return render(request, 'syllabuses/school_list.html', {'schools': schools})

    def post(self, request):
        form = SchoolForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('syllabuses:school_list')
        return render(request, 'syllabuses/school_form.html', {'form': form})

class CustomUserView(View):
    def get(self, request):
        users = CustomUser.objects.all()
        return render(request, 'syllabuses/user_list.html', {'users': users})

    def post(self, request):
        form = CustomUserForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('syllabuses:user_list')
        return render(request, 'syllabuses/user_form.html', {'form': form})

class DirectorView(View):
    def get(self, request):
        directors = Director.objects.all()
        return render(request, 'syllabuses/director_list.html', {'directors': directors})

    def post(self, request):
        form = DirectorForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('syllabuses:director_list')
        return render(request, 'syllabuses/director_form.html', {'form': form})